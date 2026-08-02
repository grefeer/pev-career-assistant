"""Attachment generation service for approved resume versions.

Produces PDF and DOCX attachments from approved facts and diffs,
then stores them as encrypted objects.
"""

from __future__ import annotations

import hashlib
import io
import json
import logging
import uuid
from typing import Any

from sqlalchemy.orm import Session

from backend.app.db.models import (
    ApprovedResumeAttachment,
)
from backend.app.services.storage import EncryptedObjectStore

logger = logging.getLogger(__name__)

# -- Rendering -----------------------------------------------------------------

TEXT_FORMATS: dict[str, str] = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

DIFF_SECTION_HEADINGS: dict[str, str] = {
    "education": "Education",
    "skills": "Skills",
    "work_experience": "Work Experience",
    "projects": "Projects",
    "awards": "Awards",
    "certifications": "Certifications",
    "summary": "Professional Summary",
}


def render_resume_lines(
    approved_facts: dict[str, Any],
    diffs: list[dict[str, Any]] | None = None,
) -> list[str]:
    """Build a deterministic list of text lines from approved facts and diffs.

    The output can be fed directly to the PDF or DOCX generators. Diffs are
    applied on top of the base facts to produce the final text.
    """
    lines: list[str] = []

    # --- Name / header -------------------------------------------------------
    name = approved_facts.get("name")
    if name:
        lines.append(name.upper())
        lines.append("")

    contact = approved_facts.get("contact")
    if contact:
        parts: list[str] = []
        for field in ("email", "phone", "location"):
            val = contact.get(field)
            if val:
                parts.append(str(val))
        if parts:
            lines.append(" | ".join(parts))
            lines.append("")

    # --- Summary -------------------------------------------------------------
    summary = approved_facts.get("summary")

    # --- Section helper ------------------------------------------------------
    def add_section(
        section_key: str,
        entries: list[dict[str, Any]] | None,
        item_fmt: str | None = None,
    ) -> None:
        heading = DIFF_SECTION_HEADINGS.get(section_key, section_key.replace("_", " ").title())
        entries = _apply_diffs_to_section(section_key, entries, diffs or [])
        if not entries:
            return
        lines.append(heading.upper())
        lines.append("")
        for entry in entries:
            if item_fmt:
                formatted = _format_entry(entry, item_fmt)
                for f_line in formatted.split("\n"):
                    lines.append(f_line)
            else:
                for key, label in _ENTRY_LABELS.get(section_key, {}).items():
                    val = entry.get(key)
                    if val:
                        lines.append(f"{label}: {val}")
            lines.append("")

    add_section("summary", [{"text": summary}] if summary else None, "text")
    add_section(
        "education",
        approved_facts.get("education"),
        "school_degree",
    )
    add_section(
        "skills",
        approved_facts.get("skills"),
        "skill_list",
    )
    add_section(
        "work_experience",
        approved_facts.get("work_experience"),
        "experience",
    )
    add_section(
        "projects",
        approved_facts.get("projects"),
        "project",
    )
    add_section(
        "awards",
        approved_facts.get("awards"),
        "award",
    )
    add_section(
        "certifications",
        approved_facts.get("certifications"),
        "certification",
    )

    return lines


_ENTRY_LABELS: dict[str, dict[str, str]] = {
    "education": {
        "school": "School",
        "degree": "Degree",
        "field": "Field",
        "start_date": "Start",
        "end_date": "End",
    },
    "work_experience": {
        "company": "Company",
        "title": "Title",
        "location": "Location",
        "start_date": "Start",
        "end_date": "End",
    },
    "projects": {
        "name": "Project",
        "role": "Role",
        "url": "URL",
    },
}

# -- Diff application ---------------------------------------------------------


def _apply_diffs_to_section(
    section_key: str,
    entries: list[dict[str, Any]] | None,
    diffs: list[dict[str, Any]],
) -> list[dict[str, Any]] | None:
    """Apply diffs that target the given section and return modified entries.

    Matching strategy (in order of precedence):
    1. Exact match of ``before`` text against the rendered entry string
    2. For ``reorder``: ``fact_ref`` interpreted as comma-separated indices
    3. For ``omit`` / ``rephrase`` / ``summarize`` / ``highlight``: ``fact_ref``
       matched against the entry's JSON hash (fallback)
    """
    if not entries:
        return entries

    section_diffs = [d for d in diffs if d.get("section") == section_key]
    if not section_diffs:
        return entries

    result = list(entries)

    for diff in section_diffs:
        op = diff.get("op")
        before = diff.get("before", "")
        fact_ref = diff.get("fact_ref", "")

        if op == "reorder":
            indices = _parse_reorder_indices(fact_ref, len(result))
            if indices:
                result = [result[i] for i in indices if i < len(result)]
            continue

        # --- Find matching entry index --------------------------------------
        match_idx = _find_entry_index(result, before, fact_ref, section_key)
        if match_idx is None:
            continue

        if op == "omit":
            result.pop(match_idx)
        elif op in ("rephrase", "summarize"):
            after = diff.get("after")
            if after is not None:
                result[match_idx] = {**result[match_idx], "_rephrased": after}
        elif op == "highlight":
            after = diff.get("after")
            if after is not None:
                result[match_idx] = {**result[match_idx], "_highlighted": after}

    return result


def _find_entry_index(
    entries: list[dict[str, Any]],
    before: str,
    fact_ref: str,
    section_key: str,
) -> int | None:
    """Locate an entry by ``before`` text, then by JSON hash, then by index."""
    # 1. Try matching by rendering the entry and comparing to ``before``
    if before:
        for i, entry in enumerate(entries):
            rendered = _format_entry(entry, _FORMAT_FOR_SECTION.get(section_key, ""))
            if before in rendered:
                return i

    # 2. Try matching by JSON hash
    if fact_ref:
        for i, entry in enumerate(entries):
            if _fact_ref(entry) == fact_ref:
                return i

    # 3. Try parsing fact_ref as a direct index
    if fact_ref and fact_ref.isdigit():
        idx = int(fact_ref)
        if 0 <= idx < len(entries):
            return idx

    return None


_FORMAT_FOR_SECTION: dict[str, str] = {
    "education": "school_degree",
    "skills": "skill_list",
    "work_experience": "experience",
    "projects": "project",
    "awards": "award",
    "certifications": "certification",
}


def _fact_ref(entry: dict[str, Any]) -> str:
    """Derive a stable fact reference from an entry (JSON-based, handles lists)."""
    return hashlib.sha256(
        json.dumps(entry, sort_keys=True, default=str).encode("utf-8")
    ).hexdigest()[:16]


def _parse_reorder_indices(raw: str, max_len: int) -> list[int]:
    """Parse a comma-separated list of indices."""
    try:
        indices = [int(x.strip()) for x in raw.split(",") if x.strip().isdigit()]
        return [i for i in indices if 0 <= i < max_len]
    except (ValueError, TypeError):
        return []


def _format_entry(entry: dict[str, Any], fmt: str) -> str:
    """Format a single entry based on the section type."""
    if fmt == "school_degree":
        school = entry.get("school", "")
        degree = entry.get("degree", "")
        field = entry.get("field", "")
        dates = _format_dates(entry)
        parts = [p for p in [degree, field, school] if p]
        line = f"{' - '.join(parts)}{dates}"
        rephrased = entry.get("_rephrased")
        return rephrased if rephrased else line

    elif fmt == "skill_list":
        skills = entry.get("skills", entry.get("name", ""))
        if isinstance(skills, list):
            skills = ", ".join(skills)
        level = entry.get("level", "")
        return f"{skills} ({level})" if level else str(skills)

    elif fmt == "experience":
        company = entry.get("company", "")
        title = entry.get("title", "")
        dates = _format_dates(entry)
        lines = [f"{title} at {company}{dates}"]
        highlights = entry.get("highlights", [])
        if isinstance(highlights, list):
            for h in highlights:
                lines.append(f"  - {h}")
        description = entry.get("description", "")
        if description:
            lines.append(f"  {description}")
        rephrased = entry.get("_rephrased")
        if rephrased:
            lines.append(f"  [{rephrased}]")
        return "\n".join(lines)

    elif fmt == "project":
        name = entry.get("name", "")
        role = entry.get("role", "")
        url = entry.get("url", "")
        parts = [f"Project: {name}"]
        if role:
            parts.append(f"Role: {role}")
        if url:
            parts.append(f"URL: {url}")
        return " | ".join(parts)

    elif fmt == "award":
        name = entry.get("name", "")
        issuer = entry.get("issuer", "")
        date = entry.get("date", "")
        parts = [f"Award: {name}"]
        if issuer:
            parts.append(issuer)
        if date:
            parts.append(date)
        return " - ".join(parts)

    elif fmt == "certification":
        name = entry.get("name", "")
        issuer = entry.get("issuer", "")
        return f"{name} ({issuer})" if issuer else name

    elif fmt == "text":
        return str(entry.get("text", ""))

    return str(entry)


def _format_dates(entry: dict[str, Any]) -> str:
    """Format start/end date range for an entry."""
    start = entry.get("start_date", "")
    end = entry.get("end_date", "")
    if start or end:
        return f" ({start} - {end})"
    return ""


# -- PDF generation (reportlab) -----------------------------------------------


def generate_resume_pdf(
    approved_facts: dict[str, Any],
    diffs: list[dict[str, Any]] | None = None,
) -> bytes:
    """Generate a PDF resume as bytes.

    Returns bytes that MUST start with ``b'%PDF-'``.
    """
    lines = render_resume_lines(approved_facts, diffs)

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()
    heading_style = ParagraphStyle(
        "ResumeHeading",
        parent=styles["Heading1"],
        fontSize=16,
        leading=20,
        spaceAfter=6,
        spaceBefore=12,
    )
    section_style = ParagraphStyle(
        "ResumeSection",
        parent=styles["Heading2"],
        fontSize=12,
        leading=15,
        spaceAfter=4,
        spaceBefore=10,
        textColor="#333333",
    )
    body_style = ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=2,
        spaceBefore=0,
    )

    story: list[Any] = []

    for line in lines:
        if not line:
            story.append(Spacer(1, 3 * mm))
        elif line.isupper() and len(line) > 2:
            # Detect section headings (all-caps lines)
            story.append(Paragraph(line.replace("\n", "<br/>"), section_style))
        elif line == lines[0] if lines else False:
            # First line (name) in heading style
            story.append(Paragraph(line.replace("\n", "<br/>"), heading_style))
        else:
            story.append(Paragraph(line.replace("\n", "<br/>"), body_style))

    doc.build(story)
    return buf.getvalue()


# -- DOCX generation (python-docx) --------------------------------------------


def generate_resume_docx(
    approved_facts: dict[str, Any],
    diffs: list[dict[str, Any]] | None = None,
) -> bytes:
    """Generate a DOCX resume as bytes.

    Returns bytes that MUST start with ``b'PK'`` (the ZIP magic number for
    Office Open XML).
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lines = render_resume_lines(approved_facts, diffs)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for i, line in enumerate(lines):
        if not line:
            # Blank line — add empty paragraph with small spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
        elif i == 0 and len(line) > 2:
            # First line = name
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.isupper() and len(line) > 2:
            # Section heading
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -- Storage helpers -----------------------------------------------------------


def generate_and_store_attachments(
    db: Session,
    user_id: str,
    draft_id: str,
    approved_resume_version_id: str,
    approved_facts: dict[str, Any],
    diffs: list[dict[str, Any]],
    object_store: EncryptedObjectStore,
) -> list[ApprovedResumeAttachment]:
    """Generate PDF and DOCX attachments, encrypt, store, and persist records."""
    attachments: list[ApprovedResumeAttachment] = []
    generators = {
        "pdf": generate_resume_pdf,
        "docx": generate_resume_docx,
    }

    for fmt, generator_fn in generators.items():
        body = generator_fn(approved_facts, diffs)
        attachment_id = str(uuid.uuid4())
        object_key = f"resumes/{user_id}/{draft_id}/{attachment_id}.{fmt}"
        content_type = TEXT_FORMATS[fmt]

        stored = object_store.put(
            key=object_key,
            plaintext=body,
            content_type=content_type,
        )

        attachment = ApprovedResumeAttachment(
            id=attachment_id,
            draft_id=draft_id,
            approved_resume_version_id=approved_resume_version_id,
            user_id=user_id,
            format=fmt,
            object_key=object_key,
            content_type=content_type,
            plaintext_size=stored.plaintext_size,
            encryption_version=stored.encryption,
            status="ready",
        )
        db.add(attachment)
        attachments.append(attachment)

    db.commit()
    return attachments


def compensate_attachments(
    db: Session,
    attachment_ids: list[str],
    object_store: EncryptedObjectStore,
) -> None:
    """Delete stored objects and mark attachment records as failed."""
    attachments = (
        db.query(ApprovedResumeAttachment)
        .filter(ApprovedResumeAttachment.id.in_(attachment_ids))
        .all()
    )
    for att in attachments:
        try:
            object_store.delete(att.object_key)
        except Exception:
            logger.warning("failed to delete object %s during compensation", att.object_key)
        att.status = "failed"
    db.commit()


def download_attachment(
    db: Session,
    attachment_id: str,
    user_id: str,
    object_store: EncryptedObjectStore,
) -> tuple[bytes, str, str]:
    """Retrieve and decrypt an attachment.

    Returns ``(body, content_type, filename)``. Raises ``PermissionError`` if
    the attachment does not belong to the requesting user.
    """
    attachment = (
        db.query(ApprovedResumeAttachment)
        .filter(ApprovedResumeAttachment.id == attachment_id)
        .first()
    )
    if attachment is None:
        raise FileNotFoundError(f"Attachment {attachment_id} not found")
    if attachment.user_id != user_id:
        raise PermissionError("Attachment does not belong to this user")

    body = object_store.get(key=attachment.object_key)
    filename = f"resume.{attachment.format}"
    return body, attachment.content_type, filename
