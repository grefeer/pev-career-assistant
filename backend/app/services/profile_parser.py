from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from backend.app.domain.profiles import (
    EvidenceCandidate,
    ParsedResumeDocument,
    UnsupportedResumeTypeError,
)


SECTION_ALIASES: dict[str, str] = {
    "教育经历": "education",
    "教育背景": "education",
    "实习经历": "experience",
    "工作经历": "experience",
    "项目经历": "projects",
    "技能": "skills",
    "专业技能": "skills",
    "获奖": "awards",
    "荣誉奖项": "awards",
    "证书": "certificates",
    "语言成绩": "languages",
    "作品链接": "portfolio_links",
}


def extract_resume_document(filename: str, raw: bytes) -> ParsedResumeDocument:
    suffix = Path(filename).suffix.lower()
    if suffix not in {".txt", ".md", ".pdf", ".docx"}:
        raise UnsupportedResumeTypeError(suffix)
    try:
        if suffix in {".txt", ".md"}:
            text = raw.decode("utf-8")
        elif suffix == ".pdf":
            reader = PdfReader(BytesIO(raw))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            document = Document(BytesIO(raw))
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    except Exception:
        return ParsedResumeDocument("", True, "resume_parse_unreadable")
    normalized = "\n".join(
        line.strip() for line in text.splitlines() if line.strip()
    )
    if not normalized:
        return ParsedResumeDocument("", True, "resume_text_unavailable")
    return ParsedResumeDocument(normalized, False, None)


def extract_evidence_candidates(text: str) -> list[EvidenceCandidate]:
    candidates: list[EvidenceCandidate] = []
    lines = text.splitlines()
    current_section: str | None = None
    section_lines: list[str] = []
    seen_skills: list[str] = []
    seen_name = False
    seen_email = False
    seen_phone = False
    _EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
    _PHONE_RE = re.compile(r"1[3-9]\d{9}")
    _URL_RE = re.compile(r"https?://\S+")

    portfolio_urls: list[str] = []

    def _add_basic_candidate(
        field_path: str,
        value: object,
        excerpt: str,
        confidence: int,
    ) -> None:
        candidates.append(
            EvidenceCandidate(
                field_path=field_path,
                candidate_value=value,
                evidence_excerpt=excerpt[:500],
                confidence=confidence,
            )
        )

    def _flush_section() -> None:
        nonlocal current_section, section_lines, seen_skills
        if current_section is None or not section_lines:
            return
        if current_section == "skills":
            items: list[str] = []
            for line in section_lines:
                for sep in ("、", ",", "，", ";", "；", " "):
                    if sep in line:
                        parts = [p.strip() for p in line.split(sep) if p.strip()]
                        items.extend(parts)
                        break
                else:
                    items.append(line.strip())
            deduped = list(dict.fromkeys(items))
            seen_skills.extend(deduped)
            _add_basic_candidate("skills", deduped, " ".join(section_lines), 85)
        else:
            # ``current_section`` is only ever assigned a value from
            # ``SECTION_ALIASES`` (line above), and every alias maps to a member
            # of ``STANDARD_FIELD_PATHS`` (or ``skills`` handled above), so this
            # branch is exhaustive for non-skills sections without a dead arm.
            cleaned = [line.strip() for line in section_lines if line.strip()]
            _add_basic_candidate(
                current_section,
                cleaned,
                " ".join(cleaned),
                80,
            )
        section_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Check for section headings
        heading_key = stripped.rstrip("：:")
        if heading_key in SECTION_ALIASES:
            _flush_section()
            current_section = SECTION_ALIASES[heading_key]
            continue

        # basics.name: first non-heading line
        if not seen_name and current_section is None and len(stripped) <= 120:
            _add_basic_candidate("basics.name", stripped, stripped, 65)
            seen_name = True
            continue

        # basics.email
        email_match = _EMAIL_RE.search(stripped)
        if not seen_email and email_match:
            _add_basic_candidate(
                "basics.email", email_match.group(0), email_match.group(0), 90
            )
            seen_email = True

        # basics.phone
        phone_match = _PHONE_RE.search(stripped)
        if not seen_phone and phone_match:
            _add_basic_candidate(
                "basics.phone", phone_match.group(0), phone_match.group(0), 90
            )
            seen_phone = True

        # portfolio_links - extract URLs even without heading
        for url in _URL_RE.findall(stripped):
            if url not in portfolio_urls:
                portfolio_urls.append(url)

        # Accumulate section content
        if current_section is not None:
            section_lines.append(stripped)
        elif seen_name:
            # After name, if no section heading yet, still collect
            pass

    _flush_section()

    # Add portfolio links found during scan
    if portfolio_urls:
        _add_basic_candidate(
            "portfolio_links",
            portfolio_urls,
            " ".join(portfolio_urls),
            80,
        )

    return candidates
